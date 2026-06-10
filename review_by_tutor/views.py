import logging
import mimetypes
from pathlib import PurePosixPath
from io import BytesIO
from urllib.parse import unquote, urlparse
from zipfile import ZIP_DEFLATED, ZipFile
from datetime import timedelta

import requests
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.mail import send_mail
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db import transaction
from django.db.models import Q, Subquery, OuterRef, Count, Exists, CharField, Value, When, Case, IntegerField
from django.http import Http404, FileResponse, JsonResponse, StreamingHttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.utils import timezone
from django.utils.encoding import smart_str
from django.utils.text import slugify
from django.views.decorators.http import require_POST, require_GET
from docx import Document as WordDocument

from Putevka import settings
from config import BASE_URL
from core.bot import send_tg_notification_to_user
from core.decorators import ensure_registration_gate
from core.forms import SendNotificationForm
from core.models import MotivationLetter, Notification, UserNotification
from core.services.email_service import send_email_to_user
from documents.models import Document
from my_study.models import CourseSelection, UniversityPriority, AssessmentResult, School, Course
from review_by_tutor.forms import MotivationLetterStaffForm, UserInfoStaffForm, ScholarVideoStaffForm, \
    DocumentStaffUploadForm, DocumentCommentForm, \
    DocumentStatusForm, InterviewForm, TestAssignmentCreateForm, TestAssignmentEditForm, TestResultForm, \
    LetterRevisionForm, MotivationLetterRubricReviewStaffForm, LetterDeadlineForm, ScholarVideoDeadlineForm, \
    ScholarVideoYandexPublicLinkForm, InterviewResultForm, TestRevisionForm
from review_by_tutor.models import Interview, TestAssignment, InterviewPreparation, InterviewTemplate, InterviewResult, \
    TestTemplate, TestingInstruction
from review_by_tutor.services.interview_xlsx import build_prefilled_interview_xlsx, import_interview_result_xlsx
from review_by_tutor.services.staff_users import build_staff_users_queryset, get_staff_users_filters
from review_by_tutor.utils.contact_form import handle_send_notification
from review_by_tutor.utils.selection_stages import require_selection_step
from scholar_form.models import UserInfo, ScholarVideo, StaffNote, InterviewInstruction
from scholar_form.services.yandex_disk import (
    YandexDiskError,
    build_public_resource_ref,
    get_download_url,
    get_public_download_url,
    get_public_resource_metadata,
    get_resource_metadata,
)
from scholar_form.views import build_video_asset_context

logger = logging.getLogger(__name__)
User = get_user_model()
VIDEO_ALLOWED_MIME = {"video/mp4", "video/webm", "video/quicktime", "video/x-quicktime", "video/x-matroska"}
VIDEO_ALLOWED_EXT = {".mp4", ".webm", ".mov", ".mkv"}
VIDEO_STREAM_CHUNK_SIZE = 1024 * 1024


def _staff_check(user):
    return user.is_staff


def _is_yandex_disk_public_url(value: str) -> bool:
    parsed = urlparse(value)
    host = (parsed.netloc or "").lower()
    return parsed.scheme in {"http", "https"} and host in {"disk.yandex.ru", "disk.yandex.com", "yadi.sk"}


def _path_from_yandex_client_url(value: str) -> str:
    parsed = urlparse(value)
    host = (parsed.netloc or "").lower()
    if host not in {"disk.yandex.ru", "disk.yandex.com"}:
        return ""

    marker = "/client/disk/"
    if marker not in parsed.path:
        return ""

    tail = parsed.path.split(marker, 1)[1].strip("/")
    return f"disk:/{unquote(tail)}" if tail else ""


def _normalize_interview_video_source(value: str) -> tuple[str, str]:
    raw = (value or "").strip()
    if not raw:
        return "", ""

    client_path = _path_from_yandex_client_url(raw)
    if client_path:
        return "yandex_disk_path", client_path

    if raw.startswith("disk:/"):
        return "yandex_disk_path", raw

    parsed = urlparse(raw)
    if parsed.scheme in {"http", "https"}:
        if _is_yandex_disk_public_url(raw):
            return "yandex_public_url", raw
        raise YandexDiskError("Поддерживаются только ссылки Яндекс Диска.")

    return "yandex_disk_path", f"disk:/{raw.lstrip('/')}"


def _guess_video_mime(name: str, fallback: str = "") -> str:
    guessed, _ = mimetypes.guess_type(name or "")
    return fallback or guessed or "application/octet-stream"


def _validate_video_metadata(*, name: str, mime: str, size: int | None):
    ext = PurePosixPath(name or "").suffix.lower()
    normalized_mime = (mime or "").split(";", 1)[0].strip().lower()

    if ext not in VIDEO_ALLOWED_EXT and normalized_mime not in VIDEO_ALLOWED_MIME:
        raise YandexDiskError("Файл на Яндекс Диске должен быть видео MP4, WebM, MOV или MKV.")

    if size is not None and size <= 0:
        raise YandexDiskError("Файл на Яндекс Диске пустой.")


def _is_video_resource(item: dict) -> bool:
    if item.get("type") != "file":
        return False
    name = item.get("name") or ""
    mime = item.get("mime_type") or _guess_video_mime(name)
    ext = PurePosixPath(name).suffix.lower()
    normalized_mime = (mime or "").split(";", 1)[0].strip().lower()
    return ext in VIDEO_ALLOWED_EXT or normalized_mime in VIDEO_ALLOWED_MIME


def _embedded_items(metadata: dict) -> list[dict]:
    embedded = metadata.get("_embedded") or {}
    items = embedded.get("items") or []
    return items if isinstance(items, list) else []


def _single_video_from_folder(metadata: dict) -> dict:
    videos = [item for item in _embedded_items(metadata) if _is_video_resource(item)]
    if not videos:
        raise YandexDiskError("В папке Яндекс Диска не найдено видеофайлов.")
    if len(videos) > 1:
        raise YandexDiskError("В папке Яндекс Диска несколько видеофайлов. Укажите ссылку или путь на конкретный файл.")
    return videos[0]


def _join_disk_child_path(parent_path: str, child_name: str) -> str:
    parent = parent_path.rstrip("/")
    if parent.startswith("disk:/"):
        return f"{parent}/{child_name}"
    return f"disk:/{parent.lstrip('/')}/{child_name}"


def _apply_interview_video_link(interview: Interview, source_value: str, uploaded_by):
    source_type, source_ref = _normalize_interview_video_source(source_value)
    interview.video_yandex_disk_url = (source_value or "").strip()
    interview.video_yandex_disk_path = ""
    interview.video_source_type = source_type
    interview.video_name = ""
    interview.video_size = None
    interview.video_mime = ""
    interview.video_link_error = ""
    interview.video_link_checked_at = timezone.now()

    if not source_ref:
        return

    if source_type == "yandex_disk_path":
        metadata = get_resource_metadata(
            source_ref,
            log_context={"interview_id": interview.pk, "user_id": interview.user_id},
        )
        if metadata.get("type") == "dir":
            child = _single_video_from_folder(metadata)
            source_ref = _join_disk_child_path(source_ref, child.get("name") or "")
            metadata = child
        elif metadata.get("type") != "file":
            raise YandexDiskError("Укажите ссылку или путь на видеофайл либо папку с одним видеофайлом.")

        name = metadata.get("name") or PurePosixPath(source_ref.replace("disk:/", "", 1)).name
        mime = metadata.get("mime_type") or _guess_video_mime(name)
        size = metadata.get("size")
        _validate_video_metadata(name=name, mime=mime, size=size)
        interview.video_yandex_disk_path = source_ref
    else:
        metadata = get_public_resource_metadata(
            source_ref,
            log_context={"interview_id": interview.pk, "user_id": interview.user_id},
        )
        public_path = ""
        if metadata.get("type") == "dir":
            child = _single_video_from_folder(metadata)
            public_path = child.get("path") or child.get("name") or ""
            metadata = child
        elif metadata.get("type") != "file":
            raise YandexDiskError("Укажите ссылку на видеофайл либо папку с одним видеофайлом.")

        name = metadata.get("name") or "video"
        mime = metadata.get("mime_type") or _guess_video_mime(name)
        size = metadata.get("size")
        _validate_video_metadata(name=name, mime=mime, size=size)
        interview.video_yandex_disk_path = public_path
        href = get_public_download_url(
            source_ref,
            public_path=public_path,
            log_context={"interview_id": interview.pk, "user_id": interview.user_id},
        )
        head = requests.head(href, allow_redirects=True, timeout=30)
        if head.status_code >= 400:
            probe = requests.get(href, headers={"Range": "bytes=0-0"}, stream=True, timeout=30)
            probe.close()
            if probe.status_code >= 400:
                raise YandexDiskError("Не удалось проверить публичный видеофайл на Яндекс Диске.")
        mime = mime or head.headers.get("Content-Type") or _guess_video_mime(name)
        size_header = head.headers.get("Content-Length")
        size = size or (int(size_header) if size_header and size_header.isdigit() else None)
        _validate_video_metadata(name=name, mime=mime, size=size)

    interview.video_name = name
    interview.video_size = size
    interview.video_mime = _guess_video_mime(name, mime)
    interview.video_uploaded_by = uploaded_by
    interview.video_uploaded_at = timezone.now()
    interview.transcript_status = "PENDING"
    interview.transcript_error = ""
    interview.transcript = ""


def _interview_video_download_href(interview: Interview) -> str:
    if interview.video_source_type == "yandex_disk_path" and interview.video_yandex_disk_path:
        return get_download_url(
            interview.video_yandex_disk_path,
            log_context={"interview_id": interview.pk, "user_id": interview.user_id},
        )
    if interview.video_source_type == "yandex_public_url" and interview.video_yandex_disk_url:
        return get_public_download_url(
            interview.video_yandex_disk_url,
            public_path=interview.video_yandex_disk_path,
            log_context={"interview_id": interview.pk, "user_id": interview.user_id},
        )
    raise Http404("Видео на Яндекс Диске не настроено")


def _apply_scholar_video_public_link(video: ScholarVideo, source_value: str):
    source_type, source_ref = _normalize_interview_video_source(source_value)
    if source_type != "yandex_public_url":
        raise YandexDiskError("Укажите публичную ссылку Яндекс Диска.")

    metadata = get_public_resource_metadata(
        source_ref,
        log_context={"scholar_video_id": video.pk, "user_id": video.user_id},
    )
    public_path = ""
    if metadata.get("type") == "dir":
        child = _single_video_from_folder(metadata)
        public_path = child.get("path") or child.get("name") or ""
        metadata = child
    elif metadata.get("type") != "file":
        raise YandexDiskError("Укажите ссылку на видеофайл либо папку с одним видеофайлом.")

    name = metadata.get("name") or "video"
    mime = metadata.get("mime_type") or _guess_video_mime(name)
    size = metadata.get("size")
    _validate_video_metadata(name=name, mime=mime, size=size)

    href = get_public_download_url(
        source_ref,
        public_path=public_path,
        log_context={"scholar_video_id": video.pk, "user_id": video.user_id},
    )
    head = requests.head(href, allow_redirects=True, timeout=30)
    if head.status_code >= 400:
        probe = requests.get(href, headers={"Range": "bytes=0-0"}, stream=True, timeout=30)
        probe.close()
        if probe.status_code >= 400:
            raise YandexDiskError("Не удалось проверить публичный видеофайл на Яндекс Диске.")

    mime = mime or head.headers.get("Content-Type") or _guess_video_mime(name)
    size_header = head.headers.get("Content-Length")
    size = size or (int(size_header) if size_header and size_header.isdigit() else None)
    _validate_video_metadata(name=name, mime=mime, size=size)

    video.yandex_disk_path = build_public_resource_ref(source_ref, public_path)
    video.yandex_disk_uploaded_at = timezone.now()
    video.yandex_disk_error = ""
    video.transcript_status = "PENDING"
    video.transcript_error = ""
    video.transcript_text = ""
    video.file = None


def _letter_display_name(user) -> str:
    full_name = (user.get_full_name() or "").strip()
    return full_name or user.username or f"user_{user.pk}"


def _safe_letter_basename(user) -> str:
    base = slugify(_letter_display_name(user), allow_unicode=True).strip("-_")
    return base or f"user_{user.pk}"


def _build_letter_docx(letter: MotivationLetter) -> bytes:
    document = WordDocument()

    document.add_heading("Мотивационное письмо", level=1)
    document.add_paragraph(f"Участник: {_letter_display_name(letter.user)}")
    document.add_paragraph(f"ID пользователя: {letter.user_id}")

    if letter.submitted_at:
        submitted_at = timezone.localtime(letter.submitted_at).strftime("%d.%m.%Y %H:%M")
        document.add_paragraph(f"Отправлено: {submitted_at}")

    text = (letter.letter_text or "").strip()
    if text:
        for block in text.splitlines():
            document.add_paragraph(block)
    else:
        document.add_paragraph("Текст письма отсутствует.")

    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def _get_downloadable_letter_or_404(user_id: int) -> MotivationLetter:
    letter = get_object_or_404(
        MotivationLetter.objects.select_related("user"),
        user_id=user_id,
    )
    if not (letter.letter_text or "").strip():
        raise Http404("Текст мотивационного письма не найден")
    return letter


@login_required
@user_passes_test(_staff_check)
@transaction.atomic
def staff_letter_detail(request, user_id: int):
    user = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user)

    letter = (
        MotivationLetter.objects.select_related("user", "rubric_review")
        .filter(user_id=user_id)
        .first()
    )

    if letter is None:
        letter = MotivationLetter.objects.create(user_id=user_id)

    revision_form = LetterRevisionForm(request.POST or None)
    deadline_form = LetterDeadlineForm(request.POST or None)

    rubric_review = getattr(letter, "rubric_review", None) if letter else None
    rubric_form = MotivationLetterRubricReviewStaffForm(
        request.POST or None,
        instance=rubric_review
    ) if rubric_review else None

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user)
        logger.info("SADF")
        return redirect("staff_letter_detail", user_id=user_id)

    if request.method == "POST":
        if letter is None:
            messages.error(request, "У пользователя ещё нет мотивационного письма — сохранять нечего.")
            return redirect("staff_letter_detail", user_id=user_id)

        if "action_toggle_favorite" in request.POST:
            letter.is_favorite = not letter.is_favorite
            letter.save(update_fields=["is_favorite", "updated_at"])
            messages.success(
                request,
                "Письмо добавлено в избранное." if letter.is_favorite else "Письмо убрано из избранного."
            )
            return redirect("staff_letter_detail", user_id=user_id)

        if "action_rubric_save" in request.POST:
            if rubric_review is None:
                messages.error(request, "Нет авторазбора по рубрике — редактировать нечего.")
                return redirect("staff_letter_detail", user_id=user_id)

            if rubric_form and rubric_form.is_valid():
                saved = rubric_form.save(commit=False)
                saved.save()
                messages.success(request, "Рубрика сохранена.")
                return redirect("staff_letter_detail", user_id=user_id)
            else:
                messages.error(request, "Исправьте ошибки в форме рубрики.")

        elif "action_revision" in request.POST:
            if revision_form.is_valid():
                letter.is_done = False
                comment = revision_form.cleaned_data["revision_comment"].strip()
                letter.status = MotivationLetter.Status.REVISION
                letter.revision_comment = comment
                letter.revision_requested_at = timezone.now()
                letter.revision_requested_by = request.user
                letter.is_done = False
                letter.save(update_fields=[
                    "status", "revision_comment", "revision_requested_at",
                    "revision_requested_by", "is_done", "updated_at",
                ])

                notify_participant(
                    user,
                    "Мотивационное письмо отправлено на доработку",
                    f"Твоё мотивационное письмо нужно доработать.\n\nКомментарий:\n{comment}",
                    sender=request.user
                )

                messages.success(request, "Письмо отправлено на дописывание.")
                return redirect("staff_letter_detail", user_id=user_id)
            else:
                messages.error(request, "Укажите комментарий для доработки.")

        elif "action_deadline_save" in request.POST:
            if deadline_form.is_valid():
                letter.deadline_at = deadline_form.cleaned_data["deadline_at"]
                letter.save(update_fields=["deadline_at"])

                deadline_str = timezone.localtime(letter.deadline_at).strftime("%d.%m.%Y %H:%M")

                notify_participant(
                    user,
                    "Установлен дедлайн по мотивационному письму",
                    f"Для мотивационного письма установлен дедлайн: {deadline_str}.",
                    sender=request.user
                )

                messages.success(request, "Дедлайн обновлён")
                return redirect(request.path)
            else:
                messages.error(request, "Укажите корректный дедлайн.")

        elif "action_deadline_clear" in request.POST:
            letter.deadline_at = None
            letter.save(update_fields=["deadline_at"])
            messages.success(request, "Дедлайна больше нет.")

        else:
            form = MotivationLetterStaffForm(request.POST, instance=letter)
            if form.is_valid():
                updated = form.save(commit=False)
                updated.save()

                letter.status = MotivationLetter.Status.SUBMITTED

                letter.save(update_fields=["status", "updated_at"])

                notify_participant(
                    user,
                    "Мотивационное письмо принято",
                    "Твоё мотивационное письмо проверено и принято.",
                    sender=request.user
                )

                messages.success(request, "Оценка/фидбэк сохранены.")
                return redirect("staff_letter_detail", user_id=user_id)
            else:
                messages.error(request, "Исправьте ошибки в форме.")

    else:
        form = MotivationLetterStaffForm(instance=letter) if letter else None

    readonly_ctx = {
        "status": getattr(letter, "status", None),
        "submitted_at": getattr(letter, "submitted_at", None),
        "revision_comment": getattr(letter, "revision_comment", None),
        "revision_requested_at": getattr(letter, "revision_requested_at", None),
    }

    ctx = {
        "user_obj": user,
        "letter": letter,
        "form": form,
        "revision_form": revision_form,
        "deadline_form": deadline_form,

        "rubric_review": rubric_review,
        "rubric_form": rubric_form,

        "active": "motivation_letter",
        "readonly": readonly_ctx,

        "send_notification_form": send_notification_form,
    }
    return render(request, "staff_templates/letter_detail.html", ctx)


@login_required
@user_passes_test(_staff_check)
@require_GET
def staff_letter_download(request, user_id: int):
    letter = _get_downloadable_letter_or_404(user_id)
    filename = f"motivation-letter_{_safe_letter_basename(letter.user)}_id{letter.user_id}.docx"
    return FileResponse(
        BytesIO(_build_letter_docx(letter)),
        as_attachment=True,
        filename=smart_str(filename),
    )


@login_required
@user_passes_test(_staff_check)
@transaction.atomic
def staff_profile_detail(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    profile = get_object_or_404(
        UserInfo.objects.select_related("user"),
        user_id=user_id
    )

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user_obj)
        return redirect("staff_profile_detail", user_id=user_id)

    if request.method == "POST":
        form = UserInfoStaffForm(request.POST, request.FILES, instance=profile)

        if form.is_valid():
            if "action_save" in request.POST:
                form.save()
                messages.success(request, "Изменения сохранены.")
                return redirect("staff_profile_detail", user_id=user_id)

            if "action_revision" in request.POST:
                obj = form.save(commit=False)
                comment = (obj.revision_comment or "").strip()

                if not comment:
                    messages.error(request, "Нужно указать причину доработки (revision_comment).")
                    return redirect("staff_profile_detail", user_id=user_id)

                obj.form_status = "revision"
                obj.revision_requested_at = timezone.now()
                obj.save()
                form.save_m2m()

                def _notify():
                    try:
                        send_tg_notification_to_user(
                            obj.user,
                            (
                                "✏️ Анкета отправлена на доработку.\n\n"
                                f"Причина:\n{comment}\n\n"
                                "Открой анкету и внеси правки."
                            ),
                            url=f"{BASE_URL}form/apply/",
                            button_text="📝 Открыть анкету"
                        )
                    except Exception as e:
                        logger.exception("Ошибка отправки Telegram уведомления (revision): %s", e)

                    user_email = obj.user.email or obj.email

                    if user_email:
                        try:
                            send_mail(
                                subject="Анкета отправлена на доработку",
                                message=(
                                    "Здравствуйте!\n\n"
                                    "Ваша анкета отправлена на доработку.\n\n"
                                    f"Причина:\n{comment}\n\n"
                                    f"Перейдите по ссылке для редактирования:\n{BASE_URL}form/apply/\n\n"
                                    "С уважением,\nКоманда программы"
                                ),
                                from_email=settings.DEFAULT_FROM_EMAIL,
                                recipient_list=[user_email],
                                fail_silently=False,
                            )
                        except Exception as e:
                            logger.exception("Ошибка отправки Email уведомления (revision): %s", e)

                transaction.on_commit(_notify)

                messages.success(
                    request,
                    "Анкета отправлена на дописывание. Пользователь будет уведомлён."
                )
                return redirect("staff_profile_detail", user_id=user_id)

            if "action_approve" in request.POST:
                obj = form.save(commit=False)
                obj.form_status = "approved"
                obj.save()
                form.save_m2m()

                def _notify():
                    message_text = (
                        "🎉 Анкета успешно принята!\n\n"
                        "Поздравляем! Твоя анкета прошла проверку.\n\n"
                        "Ожидайте дальнейшей информации."
                    )

                    try:
                        send_tg_notification_to_user(
                            obj.user,
                            message_text,
                            url=f"{BASE_URL}form/apply/",
                            button_text="📄 Открыть анкету"
                        )
                    except Exception as e:
                        logger.exception("Ошибка отправки Telegram уведомления (approve): %s", e)

                    user_email = obj.user.email or obj.email

                    if user_email:
                        try:
                            send_mail(
                                subject="Ваша анкета принята 🎉",
                                message=(
                                    "Здравствуйте!\n\n"
                                    "Ваша анкета успешно прошла проверку и принята.\n\n"
                                    "Мы свяжемся с вами для дальнейших шагов.\n\n"
                                    f"Просмотреть анкету можно здесь:\n{BASE_URL}form/apply/\n\n"
                                    "С уважением,\nКоманда программы"
                                ),
                                from_email=settings.DEFAULT_FROM_EMAIL,
                                recipient_list=[user_email],
                                fail_silently=False,
                            )
                        except Exception as e:
                            logger.exception("Ошибка отправки Email уведомления (approve): %s", e)

                transaction.on_commit(_notify)

                messages.success(request, "Анкета принята. Пользователь будет уведомлён.")
                return redirect("staff_profile_detail", user_id=user_id)

            if "action_reject" in request.POST:
                obj = form.save(commit=False)
                comment = (obj.revision_comment or "").strip()

                obj.form_status = "rejected"
                obj.save()
                form.save_m2m()

                def _notify():
                    message_text = "Анкета отклонена."
                    if comment:
                        message_text += f"\n\nКомментарий:\n{comment}"

                    try:
                        send_tg_notification_to_user(
                            obj.user,
                            message_text,
                            url=f"{BASE_URL}form/apply/",
                            button_text="Открыть анкету"
                        )
                    except Exception as e:
                        logger.exception("Ошибка отправки Telegram уведомления (reject): %s", e)

                    user_email = obj.user.email or obj.email

                    if user_email:
                        try:
                            email_message = "Здравствуйте!\n\nВаша анкета отклонена."
                            if comment:
                                email_message += f"\n\nКомментарий:\n{comment}"
                            email_message += "\n\nС уважением,\nКоманда программы"
                            send_mail(
                                subject="Ваша анкета отклонена",
                                message=email_message,
                                from_email=settings.DEFAULT_FROM_EMAIL,
                                recipient_list=[user_email],
                                fail_silently=False,
                            )
                        except Exception as e:
                            logger.exception("Ошибка отправки Email уведомления (reject): %s", e)

                transaction.on_commit(_notify)

                messages.success(request, "Анкета отклонена. Пользователь будет уведомлён.")
                return redirect("staff_profile_detail", user_id=user_id)

            form.save()
            messages.success(request, "Сохранено.")
            return redirect("staff_profile_detail", user_id=user_id)

        messages.error(request, "Исправьте ошибки в форме.")
    else:
        form = UserInfoStaffForm(instance=profile)

    send_notification_form = SendNotificationForm()

    return render(request, "staff_templates/profile_details.html", {
        "profile": profile,
        "user_obj": user_obj,
        "form": form,
        "active": "apply",
        "send_notification_form": send_notification_form,
    })


@login_required
@user_passes_test(_staff_check)
@transaction.atomic
def staff_video_detail(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user_obj)

    video = (ScholarVideo.objects
             .select_related("user")
             .filter(user_id=user_id)
             .first())

    staff_form = ScholarVideoStaffForm(instance=video) if video else None
    deadline_form = ScholarVideoDeadlineForm(instance=video)
    public_link_form = ScholarVideoYandexPublicLinkForm()

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user_obj)
        logger.info("SADF")
        return redirect("staff_video_detail", user_id=user_id)

    if request.method == "POST":
        if "action_deadline_save" in request.POST:
            if not video:
                video = ScholarVideo.objects.create(user=user_obj)
            deadline_form = ScholarVideoDeadlineForm(request.POST, instance=video)
            if deadline_form.is_valid():
                deadline_form.save()
                messages.success(request, "Дедлайн обновлён.")
                return redirect("staff_video_detail", user_id=user_id)
            messages.error(request, "Исправьте ошибки в форме дедлайна.")

        elif "action_deadline_clear" in request.POST:
            if video:
                video.deadline_at = None
                video.save(update_fields=["deadline_at"])
                messages.success(request, "Дедлайн удалён.")
            else:
                messages.info(request, "Дедлайн не задан — удалять нечего.")
            return redirect("staff_video_detail", user_id=user_id)

        elif "action_video_link_save" in request.POST:
            if not video:
                video = ScholarVideo.objects.create(user=user_obj)
            public_link_form = ScholarVideoYandexPublicLinkForm(request.POST)
            if public_link_form.is_valid():
                try:
                    _apply_scholar_video_public_link(video, public_link_form.cleaned_data["yandex_public_url"])
                except YandexDiskError as exc:
                    public_link_form.add_error("yandex_public_url", str(exc))
                    messages.error(request, str(exc))
                else:
                    video.save(update_fields=[
                        "file",
                        "yandex_disk_path",
                        "yandex_disk_uploaded_at",
                        "yandex_disk_error",
                        "transcript_status",
                        "transcript_error",
                        "transcript_text",
                        "updated_at",
                    ])
                    messages.success(request, "Видеовизитка из публичной ссылки Яндекс Диска сохранена.")
                    return redirect("staff_video_detail", user_id=user_id)
            else:
                messages.error(request, "Укажите корректную публичную ссылку Яндекс Диска.")

        else:
            staff_form = ScholarVideoStaffForm(request.POST, instance=video)
            if staff_form.is_valid():
                staff_form.save()
                messages.success(request, "Отзыв/оценка по видео сохранены.")
                logger.info("Staff %s updated ScholarVideo for user_id=%s", request.user.pk, user_id)
                return redirect("staff_video_detail", user_id=user_id)
            messages.error(request, "Исправьте ошибки в форме.")

    asset_context = build_video_asset_context(video) if video else {
        "video_download_url": None,
        "schedule_download_url": None,
        "video_name": "",
        "schedule_name": "",
        "video_mime": None,
    }

    return render(request, "staff_templates/video_detail.html", {
        "user_obj": user_obj,
        "video": video,
        "form": staff_form,
        "deadline_form": deadline_form,
        "public_link_form": public_link_form,
        "active": "my_video_page",
        "send_notification_form": send_notification_form,
        **asset_context,
    })


@login_required
@user_passes_test(_staff_check)
@transaction.atomic
def staff_documents_detail(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user_obj)

    docs = (Document.objects
            .filter(user_id=user_id)
            .prefetch_related("related_documents")
            .order_by("-uploaded_at", "-id"))

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user_obj)
        return redirect("staff_documents_detail", user_id=user_id)

    if request.method == "POST":
        form_type = request.POST.get("form_type")
        if form_type in {"update_status", "update_lock", "update_comment"}:
            doc_id = request.POST.get("document_id")
            doc = get_object_or_404(Document, pk=doc_id, user_id=user_id)

            if form_type == "update_status":
                prefix = f"st-{doc.pk}"
                form = DocumentStatusForm(request.POST, instance=doc, prefix=prefix)
            else:
                prefix = f"cm-{doc.pk}"
                form = DocumentCommentForm(request.POST, instance=doc, prefix=prefix)

            if form.is_valid():
                obj = form.save(commit=False)
                obj._ignore_lock_validation = True
                obj.save()
                messages.success(request, "Изменения сохранены.")
            else:
                logger.warning("Doc form errors for #%s: %s", doc.pk, form.errors)
                messages.error(request, "Исправьте ошибки в форме.")
            return redirect("staff_documents_detail", user_id=user_id)

        elif form_type == "upload_staff_document":
            upload_form = DocumentStaffUploadForm(request.POST, request.FILES)
            if upload_form.is_valid():
                new_doc = upload_form.save(commit=False)
                new_doc.user = user_obj
                new_doc.uploaded_by_staff = True
                new_doc._ignore_lock_validation = True
                new_doc.save()
                messages.success(request, "Документ загружен.")
            else:
                messages.error(request, "Не удалось загрузить документ. Проверьте форму.")
        else:
            messages.error(request, "Неизвестный тип формы.")
            return redirect("staff_documents_detail", user_id=user_id)

    rows = []
    for d in docs:
        rows.append((
            d,
            DocumentStatusForm(instance=d, prefix=f"st-{d.pk}"),
            DocumentCommentForm(instance=d, prefix=f"cm-{d.pk}"),
        ))
    upload_form = locals().get("upload_form", DocumentStaffUploadForm())

    return render(request, "staff_templates/documents_detail.html", {
        "user_obj": user_obj,
        "rows": rows,
        "upload_form": upload_form,
        "active": "documents_dashboard",
        "send_notification_form": send_notification_form,

    })


@login_required
@user_passes_test(_staff_check)
def staff_study_detail(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user_obj)

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user_obj)
        logger.info("SADF")
        return redirect("staff_study_detail", user_id=user_id)

    selections = (
        CourseSelection.objects
        .select_related("course__school", "course__subject")
        .filter(user_id=user_id)
        .order_by("-created_at", "-id")
    )

    priorities = (
        UniversityPriority.objects
        .filter(user_id=user_id)
        .order_by("priority", "id")
    )

    assessments = (
        AssessmentResult.objects
        .select_related("subject")
        .filter(user_id=user_id)
        .order_by("-date", "-id")
    )

    return render(request, "staff_templates/study_detail.html", {
        "user_obj": user_obj,
        "selections": selections,
        "priorities": priorities,
        "assessments": assessments,
        "active": "study",
        "send_notification_form": send_notification_form,

    })


@require_POST
@login_required
@user_passes_test(_staff_check)
def staff_note_delete(request, user_id: int, note_id: int):
    note = get_object_or_404(StaffNote, pk=note_id, target_user_id=user_id)
    note.delete()
    messages.success(request, "Заметка удалена.")
    page = (request.POST.get("page") or request.GET.get("page") or "1").strip()
    url = reverse("staff_notes", kwargs={"user_id": user_id})
    return redirect(f"{url}?page={page}")


@login_required
@user_passes_test(_staff_check)
def staff_notes_by_user(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user_obj)

    if request.method == "POST":
        text = (request.POST.get("text") or "").strip()
        if text:
            StaffNote.objects.create(target_user=user_obj, author=request.user, text=text)
            messages.success(request, "Запись добавлена.")
            return redirect("staff_notes", user_id=user_id)
        messages.error(request, "Текст записи обязателен.")

    notes_qs = (
        StaffNote.objects
        .select_related("author")
        .filter(target_user=user_obj)
        .order_by("-is_favorite", "-created_at")
    )

    page = request.GET.get("page", 1)
    paginator = Paginator(notes_qs, 5)
    try:
        notes = paginator.page(page)
    except PageNotAnInteger:
        notes = paginator.page(1)
    except EmptyPage:
        notes = paginator.page(paginator.num_pages)

    try:
        letter = MotivationLetter.objects.select_related("user").get(user=user_obj)
    except MotivationLetter.DoesNotExist:
        letter = None

    documents = (Document.objects
                 .filter(user=user_obj, is_deleted=False)
                 .order_by("-uploaded_at")[:50])

    try:
        video = ScholarVideo.objects.select_related("user").get(user=user_obj)
    except ScholarVideo.DoesNotExist:
        video = None

    ctx = {
        "user_obj": user_obj,
        "notes": notes,
        "total_count": paginator.count,
        "letter": letter,
        "documents": documents,
        "video": video,
        "active": "notes",
        "is_candidate": user_obj.user_info.status == "CANDIDATE",
        "send_notification_form": send_notification_form,

    }
    return render(request, "staff_templates/staff_notes_by_user.html", ctx)


@require_POST
@login_required
@user_passes_test(_staff_check)
def staff_note_toggle_favorite(request, user_id: int, note_id: int):
    note = get_object_or_404(StaffNote, pk=note_id, target_user_id=user_id)

    note.is_favorite = not note.is_favorite
    note.save(update_fields=["is_favorite"])

    page = (request.POST.get("page") or request.GET.get("page") or "1").strip()
    url = reverse("staff_notes", kwargs={"user_id": user_id})
    return redirect(f"{url}?page={page}")


@login_required
@user_passes_test(_staff_check)
def staff_users_list(request):
    qs = build_staff_users_queryset(request)

    paginator = Paginator(qs, 30)
    page_obj = paginator.get_page(request.GET.get("page"))

    filters_data = get_staff_users_filters(request)

    schools = School.objects.all().order_by("name")
    courses_qs = Course.objects.all()
    if filters_data["school"]:
        courses_qs = courses_qs.filter(school_id=filters_data["school"])
    courses = courses_qs.order_by("title")

    profiles = list(UserInfo.InternalStudyProfile.choices or [])
    steps = getattr(UserInfo, "SELECTION_STEP_CHOICES", None) or UserInfo._meta.get_field("selection_step").choices

    context = {
        "page_obj": page_obj,
        "schools": schools,
        "courses": courses,
        "profiles": profiles,
        "steps": steps,
        **filters_data,
    }

    return render(request, "staff_templates/users_list.html", context)


def _staff_users_queryset_from_request(request):
    return build_staff_users_queryset(request)


@login_required
@user_passes_test(_staff_check)
@require_GET
def staff_users_ids(request):
    qs = _staff_users_queryset_from_request(request)
    ids = list(qs.values_list("id", flat=True))
    return JsonResponse({"ids": ids})


@login_required
@user_passes_test(_staff_check)
@require_POST
def staff_letters_download_zip(request):
    raw_ids = request.POST.getlist("ids")
    ids = [int(x) for x in raw_ids if str(x).isdigit()]

    if not ids:
        messages.error(request, "Выберите хотя бы одного пользователя.")
        return redirect("staff_users_list")

    letters = [
        letter for letter in
        MotivationLetter.objects.select_related("user")
        .filter(user_id__in=ids)
        .exclude(letter_text__isnull=True)
        .exclude(letter_text__exact="")
        .order_by("user_id")
        if (letter.letter_text or "").strip()
    ]

    if not letters:
        messages.error(request, "У выбранных пользователей нет текстов мотивационных писем.")
        return redirect("staff_users_list")

    zip_buffer = BytesIO()
    used_names: set[str] = set()

    with ZipFile(zip_buffer, "w", compression=ZIP_DEFLATED) as archive:
        for letter in letters:
            base_name = f"motivation-letter_{_safe_letter_basename(letter.user)}_id{letter.user_id}"
            filename = f"{base_name}.docx"
            suffix = 2
            while filename in used_names:
                filename = f"{base_name}_{suffix}.docx"
                suffix += 1
            used_names.add(filename)
            archive.writestr(filename, _build_letter_docx(letter))

    zip_buffer.seek(0)
    ts = timezone.localtime().strftime("%Y%m%d_%H%M%S")
    return FileResponse(
        zip_buffer,
        as_attachment=True,
        filename=f"motivation-letters_{ts}.zip",
    )


@staff_member_required
@require_POST
def staff_send_notification(request):
    raw_ids = request.POST.getlist('ids')
    message_text = (request.POST.get('message') or '').strip()
    include_inactive = request.POST.get('include_inactive') == '1'

    ids = [int(x) for x in raw_ids if str(x).isdigit()]
    if not ids:
        messages.error(request, "Не выбраны пользователи.")
        return redirect('staff_users_list')

    if not message_text:
        messages.error(request, "Введите текст сообщения.")
        return redirect('staff_users_list')

    qs = User.objects.filter(id__in=ids).only('id', 'is_active')
    if not include_inactive:
        qs = qs.filter(is_active=True)
    final_ids = list(qs.values_list('id', flat=True))
    if not final_ids:
        messages.error(request, "Нет подходящих получателей (все неактивны или не найдены).")
        return redirect('staff_users_list')

    with transaction.atomic():
        notif = Notification.objects.create(message=message_text, sender=request.user)
        for uid in final_ids:
            link, created = UserNotification.objects.get_or_create(notification=notif, recipient_id=uid)

    logger.info(f'Оповещение {notif.pk} создано для {len(final_ids)} пользователей')
    messages.success(request, f"Оповещение «{message_text[:50]}…» отправлено {len(final_ids)} пользователям.")
    return redirect('staff_users_list')


def interview_detail(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    send_notification_form = handle_send_notification(request, user_obj)
    interview, _ = Interview.objects.get_or_create(user=user_obj)

    sections = [
        ("school", "1. Школа", [
            "school_number", "school_type", "school_distance_km", "school_distance_minutes",
            "school_specialization", "school_students_total", "school_left_after_9_est",
            "school_students_11", "class_profile", "has_ege_teachers_all",
            "teach_quality_ru", "teach_quality_math", "teach_quality_phys", "teach_quality_chem",
            "teach_quality_bio", "teach_quality_inf", "teach_quality_geo", "teach_quality_soc",
            "teach_quality_lit", "teach_quality_hist", "teach_quality_lang",
            "triples_reason", "favorite_teacher", "favorite_subject",
            "has_computer_lab", "olympiads_frequency",
            "clubs_info", "olympiad_support_by_school", "other_school_notes",
        ]),
        ("prep", "2. Необходимая подготовка", [
            "aims_medal", "admission_way", "ege_subjects", "mock_ru", "mock_math_base", "mock_math_prof",
            "mock_phys", "mock_chem", "mock_bio", "mock_inf", "mock_geo", "mock_soc", "mock_lit",
            "mock_hist", "mock_lang", "target_ru", "target_math_base", "target_math_prof", "target_phys",
            "target_chem", "target_bio", "target_inf", "target_geo", "target_soc", "target_lit", "target_hist",
            "target_lang", "had_tutor", "tutor_details", "had_online_courses", "online_courses_details",
            "olympiad_experience", "olympiads_planned", "need_olympiad_prep", "specialties", "need_career_guidance",
            "universities", "need_university_help", "why_higher_education", "prep_9_10",
            "prep_10_11", "ready_to_move", "discussed_with_parents", "other_support_needed",
        ]),
        ("family", "3. Состав семьи", [
            "family_structure", "family_many_children", "family_people_count",
            "siblings_info", "grandparents_info", "dependents_info",
            "has_disabled_need_care", "candidate_orphan", "candidate_disabled",
            "breadwinner_loss", "family_other_circumstances",
        ]),
        ("income", "4. Работа родителей, доход", [
            "mother_job", "mother_has_he",
            "father_job", "father_has_he",
            "step_parent_job", "step_parent_has_he",
            "parents_self_employed_details", "other_relatives_jobs",
            "parent_on_pension_or_care", "why_parent_not_working",
            "alimony_paid", "benefits_received", "low_income_recognized",
            "family_other_notes", "parents_involved_in_study",
            "siblings_interfere_study", "household_load",
        ]),
        ("housing", "5. Условия проживания", [
            "settlement_status", "distance_to_reg_center_km",
            "housing_type", "utilities",
            "own_room", "own_workdesk",
            "own_computer", "own_phone", "supports_whatsapp_telegram",
            "has_printer", "home_internet", "phone_internet",
            "family_has_car", "relatives_in_big_cities", "pets",
            "has_bank_card", "summer_holidays", "financial_notes",
        ]),
        ("interests", "6. Интересы и личные качества", [
            "weekday_routine", "weekend_routine",
            "clubs_hobbies", "volunteering",
            "gto_passed", "sport_info",
            "studies_extra_resources_frequency", "self_study_example",
            "other_resources", "reads_books_frequency", "favorite_book",
            "favorite_games", "favorite_movies", "favorite_socials",
            "friends_count_info", "friends_admission_discussion",
            "part_time_job", "other_achievements",
            "success_qualities", "success_definition",
            "unfinished_cases", "asks_for_help_how",
        ]),
        ("fund", "7. Работа с фондом", [
            "heard_about_fund", "parents_know_and_agree",
            "selection_experience", "knows_support_program", "most_useful_expected",
            "would_participate_without_stipend", "understands_group_courses",
            "knows_our_schools", "understands_homework_need",
            "plan_to_combine", "ready_regular_contact", "will_inform_if_absent",
            "preferred_contact_method", "ready_for_chats_webinars",
            "interesting_topics", "ready_additional_tests", "helpful_materials",
            "ready_tell_school", "ready_mentor_next",
            "fund_questions", "understands_next_steps",
        ]),
        ("final", "8. Прочее / выводы", [
            "other_notes",
            "interviewer_summary", "interviewer_risks", "interviewer_recommendations",
            "interviewer_score",
        ]),
    ]

    template_obj = (
        InterviewTemplate.objects.filter(is_active=True).order_by("-uploaded_at").first()
    )

    result_obj, _ = InterviewResult.objects.get_or_create(interview=interview)

    form = InterviewForm(instance=interview)
    result_form = InterviewResultForm(instance=result_obj)

    if request.method == "POST" and request.POST.get("action") == "send_notification":
        handle_send_notification(request, user_obj)
        logger.info("SADF")
        return redirect("interview_detail", user_id=user_id)

    if request.method == "POST":
        action = request.POST.get("op")

        if action == "save_interview_files":
            form = InterviewForm(request.POST, request.FILES, instance=interview)
            if form.is_valid():
                obj = form.save(commit=False)
                video_link = form.cleaned_data.get("video_yandex_disk_url", "")

                if "video" in request.FILES:
                    obj.video_uploaded_by = request.user
                    obj.video_uploaded_at = timezone.now()
                    obj.video_yandex_disk_url = ""
                    obj.video_yandex_disk_path = ""
                    obj.video_source_type = ""
                    obj.video_name = ""
                    obj.video_size = None
                    obj.video_mime = ""
                    obj.video_link_error = ""
                    obj.video_link_checked_at = None
                    obj.transcript_status = "PENDING"
                    obj.transcript_error = ""
                    obj.transcript = ""
                elif video_link:
                    try:
                        _apply_interview_video_link(obj, video_link, request.user)
                    except YandexDiskError as exc:
                        form.add_error("video_yandex_disk_url", str(exc))
                        messages.error(request, str(exc))
                        ctx = {
                            "user_obj": user_obj,
                            "form": form,
                            "interview": interview,
                            "template_obj": template_obj,
                            "active": "interview",
                            "result_form": result_form,
                            "interview_sections": sections,
                            "send_notification_form": send_notification_form,
                        }
                        return render(request, "staff_templates/interview_detail.html", ctx)
                else:
                    obj.video_yandex_disk_path = ""
                    obj.video_source_type = ""
                    obj.video_name = ""
                    obj.video_size = None
                    obj.video_mime = ""
                    obj.video_link_error = ""
                    obj.video_link_checked_at = None

                obj.save()
                messages.success(request, "Файлы собеседования сохранены.")
                return redirect("interview_detail", user_id=user_id)

        elif action == "upload_interview_template":
            uploaded_file = request.FILES.get("filled_template")
            if not uploaded_file:
                messages.error(request, "Выберите заполненный шаблон Excel.")
                return redirect("interview_detail", user_id=user_id)
            if not uploaded_file.name.lower().endswith(".xlsx"):
                messages.error(request, "Загрузите файл формата .xlsx.")
                return redirect("interview_detail", user_id=user_id)

            try:
                updated_fields = import_interview_result_xlsx(uploaded_file, result_obj)
            except Exception as exc:
                logger.warning(
                    "Failed to import interview result xlsx user_id=%s interview_id=%s error=%s",
                    user_id,
                    interview.pk,
                    exc,
                )
                messages.error(
                    request,
                    "Не удалось разобрать шаблон. Проверьте, что это заполненный .xlsx шаблон интервью.",
                )
                return redirect("interview_detail", user_id=user_id)

            if not updated_fields:
                messages.warning(request, "В шаблоне не найдены заполненные поля для импорта.")
            else:
                uploaded_file.seek(0)
                interview.filled_template.save(uploaded_file.name, uploaded_file, save=False)
                interview.filled_uploaded_at = timezone.now()
                interview.save(update_fields=["filled_template", "filled_uploaded_at", "updated_at"])
                messages.success(request, f"Импортировано полей: {len(updated_fields)}.")
            return redirect("interview_detail", user_id=user_id)

        elif action == "advance_selection_step":
            user_info = getattr(user_obj, "user_info", None)
            if user_info is None:
                messages.error(request, "У участника не найдена анкета, этап отбора не изменён.")
                return redirect("interview_detail", user_id=user_id)

            current_step = user_info.selection_step
            next_step = UserInfo.SelectionStep.AFTER_INTERVIEW
            if current_step == next_step:
                messages.info(request, "Участник уже находится на этапе «После собеседования».")
            else:
                user_info.selection_step = next_step
                user_info.save(update_fields=["selection_step"])
                messages.success(request, "Участник переведён на этап «После собеседования».")
            return redirect("interview_detail", user_id=user_id)

        elif action == "save_interview_result":
            result_form = InterviewResultForm(request.POST, instance=result_obj)
            if result_form.is_valid():
                result_form.save()
                messages.success(request, "Результаты интервью сохранены.")
                return redirect("interview_detail", user_id=user_id)
            else:
                messages.error(request, "Ошибка сохранения результатов интервью.")

        else:
            messages.error(request, "Неизвестное действие.")

    else:
        form = InterviewForm(instance=interview)
        result_form = InterviewResultForm(instance=result_obj)

    ctx = {
        "user_obj": user_obj,
        "form": form,
        "interview": interview,
        "template_obj": template_obj,
        "active": "interview",
        "result_form": result_form,
        "interview_sections": sections,
        "send_notification_form": send_notification_form,

    }
    return render(request, "staff_templates/interview_detail.html", ctx)


@login_required
@user_passes_test(_staff_check)
def interview_video_stream(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    interview = get_object_or_404(Interview, user=user_obj)

    if interview.video and not interview.video_source_type:
        return FileResponse(interview.video.open("rb"), content_type=interview.video_mime or None)

    href = _interview_video_download_href(interview)
    headers = {}
    range_header = request.headers.get("Range")
    if range_header:
        headers["Range"] = range_header

    try:
        upstream = requests.get(href, headers=headers, stream=True, timeout=60)
    except requests.RequestException as exc:
        logger.warning("Failed to open interview video stream user_id=%s error=%s", user_id, exc)
        raise Http404("Видео сейчас недоступно")

    if upstream.status_code >= 400:
        logger.warning(
            "Yandex Disk interview video stream failed user_id=%s status=%s body=%s",
            user_id,
            upstream.status_code,
            (upstream.text or "")[:300],
        )
        raise Http404("Видео сейчас недоступно")

    status = 206 if upstream.status_code == 206 else 200
    content_type = upstream.headers.get("Content-Type") or interview.video_mime or "application/octet-stream"
    response = StreamingHttpResponse(
        upstream.iter_content(chunk_size=VIDEO_STREAM_CHUNK_SIZE),
        status=status,
        content_type=content_type,
    )

    for header in ("Content-Length", "Content-Range", "Accept-Ranges"):
        value = upstream.headers.get(header)
        if value:
            response[header] = value
    response["Cache-Control"] = "private, no-store"
    return response


@login_required
@user_passes_test(_staff_check)
def download_interview_template(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)

    template = (
        InterviewTemplate.objects
        .filter(is_active=True)
        .order_by("-uploaded_at")
        .first()
    )
    if not template:
        raise Http404("Шаблон не найден")

    full_name = user_obj.get_full_name() or user_obj.username
    safe_name = full_name.replace(" ", "_")

    ext = template.file.name.split(".")[-1]
    filename = f"Interview_{safe_name}_ID{user_obj.id}.{ext}"

    response = FileResponse(
        template.file.open("rb"),
        as_attachment=True,
        filename=smart_str(filename),
    )
    return response


@login_required
@user_passes_test(_staff_check)
def download_prefilled_interview_template(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    interview, _ = Interview.objects.get_or_create(user=user_obj)
    result_obj, _ = InterviewResult.objects.get_or_create(interview=interview)

    try:
        payload = build_prefilled_interview_xlsx(user_obj, interview, result_obj)
    except FileNotFoundError:
        raise Http404("РЁР°Р±Р»РѕРЅ РЅРµ РЅР°Р№РґРµРЅ")

    safe_name = _safe_letter_basename(user_obj)
    filename = f"Interview_prefilled_{safe_name}_ID{user_obj.id}.xlsx"
    response = FileResponse(
        BytesIO(payload),
        as_attachment=True,
        filename=smart_str(filename),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    return response


@login_required
@user_passes_test(_staff_check)
def download_uploaded_interview_template(request, user_id: int):
    user_obj = get_object_or_404(User, pk=user_id)
    interview = get_object_or_404(Interview, user=user_obj)
    if not interview.filled_template:
        raise Http404("Загруженный шаблон интервью не найден")

    filename = PurePosixPath(interview.filled_template.name).name or f"Interview_uploaded_ID{user_obj.id}.xlsx"
    return FileResponse(
        interview.filled_template.open("rb"),
        as_attachment=True,
        filename=smart_str(filename),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )



@login_required
@require_selection_step(UserInfo.SelectionStep.TEST)
def testing_list_for_candidate(request):
    items = (TestAssignment.objects
             .filter(user=request.user)
             .order_by("-assigned_at", "-id"))
    return render(request, "testing.html", {"items": items, "user_obj": request.user, "active": "testing", "testing_instruction": TestingInstruction.get_current()})


@user_passes_test(_staff_check)
def testing_list_for_user(request, user_id):
    send_notification_form = handle_send_notification(request, User.objects.get(pk=user_id))
    items = (TestAssignment.objects
             .select_related("user", "assigned_by", "result_filled_by")
             .filter(user_id=user_id)
             .order_by("-assigned_at", "-id"))
    return render(request, "staff_templates/testing/list.html",
                  {"items": items, "target_user_id": user_id, "user_obj": get_object_or_404(User, pk=user_id),
                   "active": 'testing', "send_notification_form": send_notification_form,
})


@user_passes_test(_staff_check)
def testing_create(request):
    fixed_user_id = request.GET.get("user_id")

    fixed_user = None
    if fixed_user_id:
        fixed_user = User.objects.filter(pk=fixed_user_id).only("id", "username").first()

    if request.method == "POST":
        form = TestAssignmentCreateForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)


            if fixed_user:
                obj.user = fixed_user

                if fixed_user.user_info.selection_step == UserInfo.SelectionStep.FORM:
                    fixed_user.user_info.selection_step = UserInfo.SelectionStep.TEST
                    fixed_user.user_info.save()

            obj.assigned_by = request.user
            obj.save()

            notify_participant(
                user=obj.user,
                subject="Назначен новый тест",
                message=(
                    f"Тебе назначен тест: «{obj.title}».\n"
                    f"{'Дедлайн: ' + obj.due_at.strftime('%d.%m.%Y %H:%M') if obj.due_at else ''}\n"
                    f"Ссылка: {BASE_URL + '/form/testing'}\n"
                    f"{'Комментарий: ' + obj.instructions if obj.instructions else ''}"
                ).strip(),
                sender=request.user,
            )
            return redirect("staff_testing_list_for_user", user_id=obj.user_id)
    else:
        if fixed_user:
            form = TestAssignmentCreateForm(initial={"user": fixed_user.id})
        else:
            form = TestAssignmentCreateForm()

    return render(request, "staff_templates/testing/form.html", {
        "form": form,
        "title": "Назначить тест",
        "fixed_user": fixed_user,
    })



@user_passes_test(_staff_check)
def testing_edit(request, pk):
    obj = get_object_or_404(TestAssignment, pk=pk)
    if request.method == "POST":
        form = TestAssignmentEditForm(request.POST, instance=obj)
        if form.is_valid():
            form.save()
            return redirect("staff_testing_list_for_user", user_id=obj.user_id)
    else:
        form = TestAssignmentEditForm(instance=obj)
    return render(request, "staff_templates/testing/form.html",
                  {"form": form, "title": "Редактировать тест", "user_obj": get_object_or_404(User, pk=obj.user.id)})


def notify_participant(user, subject: str, message: str, sender=None):
    notif = Notification.objects.create(
        message=message,
        sender=sender
    )
    UserNotification.objects.create(notification=notif, recipient=user)

@user_passes_test(_staff_check)
def testing_fill_result(request, pk):
    obj = get_object_or_404(TestAssignment, pk=pk)

    res_form = TestResultForm(instance=obj)
    rev_form = TestRevisionForm(instance=obj)

    if request.method == "POST":
        action = (request.POST.get("action") or "").strip()

        if action == "revision":
            rev_form = TestRevisionForm(request.POST, instance=obj)

            if rev_form.is_valid():
                old_status = obj.status

                comment = rev_form.cleaned_data["revision_comment"]
                obj.mark_needs_revision(by_user=request.user, comment=comment)
                obj.save()

                if old_status != obj.Status.NEEDS_REVISION:
                    notify_participant(
                        user=obj.user,
                        subject="Тест отправлен на дописывание",
                        message=(
                            f"✍️ Тест «{obj.title}» отправлен на дописывание.\n\n"
                            f"Комментарий:\n{obj.revision_comment}"
                        ),
                        sender=request.user,
                    )

                return redirect("staff_testing_list_for_user", user_id=obj.user_id)


        else:
            res_form = TestResultForm(request.POST, instance=obj)
            rev_form = TestRevisionForm(instance=obj)
            if res_form.is_valid():
                filled = res_form.save(commit=False)
                filled.result_filled_by = request.user
                filled.result_filled_at = timezone.now()
                filled.mark_completed()
                filled.save()
                notify_participant(
                    user=filled.user,
                    subject="Результат теста внесён",
                    message=(
                        f"✅ По тесту «{filled.title}» внесён результат.\n\n"
                        f"{'Комментарий: ' + filled.result_text if filled.result_text else ''}"
                    ).strip(),
                    sender=request.user,
                )

                return redirect("staff_testing_list_for_user", user_id=obj.user_id)

    else:
        res_form = TestResultForm(instance=obj)
        rev_form = TestRevisionForm(instance=obj)

    return render(
        request,
        "staff_templates/testing/result_form.html",
        {
            "form": res_form,
            "rev_form": rev_form,
            "obj": obj,
            "user_obj": get_object_or_404(User, pk=obj.user_id),
            "active": "testing",
        }
    )



@ensure_registration_gate('protected')
@require_selection_step(UserInfo.SelectionStep.INTERVIEW_PREP)
@login_required
def interview_preparation_view(request):
    interview_instruction = InterviewInstruction.get_current()
    user_info = getattr(request.user, "user_info", None)
    is_after_interview = bool(
        user_info and user_info.selection_step == UserInfo.SelectionStep.AFTER_INTERVIEW
    )

    if request.method == "POST" and is_after_interview:
        user_info.after_interview_parents_notified = (
            user_info.after_interview_parents_notified
            or "after_interview_parents_notified" in request.POST
        )
        user_info.after_interview_documents_ready = (
            user_info.after_interview_documents_ready
            or "after_interview_documents_ready" in request.POST
        )
        user_info.save(update_fields=[
            "after_interview_parents_notified",
            "after_interview_documents_ready",
        ])
        messages.success(request, "Информация сохранена.")
        return redirect("preparation")

    return render(request, "interview_preparation.html", {
        "active": "interview",
        "interview_instruction": interview_instruction,
        "is_after_interview": is_after_interview,
        "user_info": user_info,
    })


@user_passes_test(_staff_check)
@require_GET
def testing_template_payload(request):
    template_id = (request.GET.get("template_id") or "").strip()
    if not template_id:
        return JsonResponse({"ok": False, "error": "template_id required"}, status=400)

    tpl = TestTemplate.objects.filter(pk=template_id, is_active=True).first()
    if not tpl:
        return JsonResponse({"ok": False, "error": "not found"}, status=404)

    due_at_str = ""
    if tpl.default_due_days:
        dt = timezone.localtime(timezone.now() + timedelta(days=tpl.default_due_days))
        due_at_str = dt.strftime("%Y-%m-%dT%H:%M")

    return JsonResponse({
        "ok": True,
        "title": tpl.title or "",
        "external_url": tpl.external_url or "",
        "instructions": tpl.instructions or "",
        "due_at": due_at_str,
    })
